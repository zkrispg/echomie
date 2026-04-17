# -*- coding: utf-8 -*-
"""Generate EchoMie-测试报告.docx from docs/EchoMie-测试报告.md (requires python-docx)."""
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Windows 常见中文字体；eastAsia 必须与西文 name 一致指向 CJK 字体，否则易出现方框
FONT_BODY = "微软雅黑"
FONT_TITLE = "黑体"
FONT_TABLE = "微软雅黑"


def _set_run_font(run, name: str, size_pt: Optional[float] = None, bold: Optional[bool] = None):
    run.font.name = name
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def format_paragraph_runs(p, font_name: str, size_pt: Optional[float] = None, bold: Optional[bool] = None):
    for run in p.runs:
        _set_run_font(run, font_name, size_pt, bold)


def set_cell_font(cell, size_pt: float = 9):
    for p in cell.paragraphs:
        for r in p.runs:
            _set_run_font(r, FONT_TABLE, size_pt)


def add_table_from_rows(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            if j < ncols:
                cell = tbl.rows[i].cells[j]
                cell.text = text.strip()
                set_cell_font(cell, 9 if i > 0 else 10)
    doc.add_paragraph()


def _is_md_table_separator(parts):
    if not parts:
        return False
    return all(p and set(p) <= set("-:") for p in parts)


def parse_md_table(lines, start_idx):
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        parts = [p.strip() for p in line.split("|")]
        parts = [p for p in parts if p != ""]
        if _is_md_table_separator(parts):
            i += 1
            continue
        rows.append(parts)
        i += 1
    return rows, i


def _strip_inline_md(s: str) -> str:
    s = s.replace("`", "")
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    return s


def _apply_doc_default_fonts(doc: Document):
    """整篇文档默认用中文字体，避免正文/标题出现方框。"""
    for name in ("Normal", "List Paragraph", "List Bullet", "List Number"):
        try:
            st = doc.styles[name]
            st.font.name = FONT_BODY
            st.font.size = Pt(10.5)
            rpr = st._element.get_or_add_rPr()
            from docx.oxml import OxmlElement

            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            for tag in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(tag), FONT_BODY)
        except KeyError:
            pass
    for lvl in (1, 2, 3):
        try:
            st = doc.styles[f"Heading {lvl}"]
            st.font.name = FONT_TITLE
            rpr = st._element.get_or_add_rPr()
            from docx.oxml import OxmlElement

            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            for tag in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(tag), FONT_TITLE)
        except KeyError:
            pass


def main():
    root = Path(__file__).resolve().parent
    md_path = root / "docs" / "EchoMie-测试报告.md"
    out_path = root / "docs" / "EchoMie-测试报告.docx"

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    _apply_doc_default_fonts(doc)

    sect = doc.sections[0]
    sect.top_margin = Cm(2)
    sect.bottom_margin = Cm(2)
    sect.left_margin = Cm(2.5)
    sect.right_margin = Cm(2.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("# "):
            p = doc.add_paragraph(_strip_inline_md(stripped[2:].strip()))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph_runs(p, FONT_TITLE, 18, True)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(_strip_inline_md(stripped[3:].strip()), level=1)
            format_paragraph_runs(p, FONT_TITLE, 14, True)
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(_strip_inline_md(stripped[4:].strip()), level=2)
            format_paragraph_runs(p, FONT_TITLE, 12, True)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            rows, ni = parse_md_table(lines, i)
            if rows:
                add_table_from_rows(doc, rows)
            i = ni
            continue

        if stripped.startswith("- "):
            body = _strip_inline_md(stripped[2:].strip())
            p = doc.add_paragraph(body, style="List Bullet")
            format_paragraph_runs(p, FONT_BODY, 10.5)
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            body = _strip_inline_md(re.sub(r"^\d+\.\s*", "", stripped))
            p = doc.add_paragraph(body, style="List Number")
            format_paragraph_runs(p, FONT_BODY, 10.5)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph(stripped.strip("*"))
            format_paragraph_runs(p, FONT_BODY, 10.5, True)
            i += 1
            continue

        if stripped:
            p = doc.add_paragraph(_strip_inline_md(stripped))
            format_paragraph_runs(p, FONT_BODY, 10.5)
        i += 1

    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    main()
